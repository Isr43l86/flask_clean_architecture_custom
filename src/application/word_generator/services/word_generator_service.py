import io
import re
from abc import ABC
from datetime import datetime
from io import BytesIO
from typing import List

import pytz
from docx import Document

from ...constants import app_constants

SECTIONS_WITH_TABLES = {
    'Criticidad': 'GARANTÍA TÉCNICA',
    'TIEMPO DEL PROYECTO': 'TIEMPO DE EJECUCIÓN DEL PROYECTO'
}
VERSION_TABLE_FLAG = 'Fecha'
PROJECT_DURATION_TABLE_FLAG = 'TIEMPO DEL PROYECTO'


def remove_numbers(input_string):
    result = re.sub(r'\d+', '', input_string)
    result = re.sub(r'\.', '', result)
    return result.strip()


def delete_sections(paragraphs, flag_style: str, sections_to_delete: List[str]):
    delete_items = False

    for paragraph in paragraphs:
        if paragraph.style.name.startswith(flag_style) and str(remove_numbers(paragraph.text)) in sections_to_delete:
            paragraph.clear()
            delete_items = True
        elif paragraph.style.name.startswith(flag_style):
            delete_items = False

        if delete_items:
            p = paragraph._element
            p.getparent().remove(p)
            p._p = p._element = None


class WordGeneratorServiceImpl(ABC):

    @staticmethod
    def generate_word_document(data: dict, path: str, sections_to_delete: List[dict]) -> tuple[BytesIO, str]:
        new_dict = {f"[{key}]": str(value).upper() for key, value in data.items()}
        new_dict['[currentDate]'] = str(datetime.now(pytz.timezone(app_constants.APP_TIMEZONE)).strftime('%m-%d-%Y'))
        new_dict['[projectVersion]'] = '1'

        doc = Document(path)

        primary_sections_to_delete = [x['sectionName'] for x in sections_to_delete if x['isSection']]
        secondary_sections_to_delete = [x['sectionName'] for x in sections_to_delete if not x['isSection']]

        for activeTable in doc.tables:

            first_row_column_text = activeTable.cell(0, 0).paragraphs[0].text

            if first_row_column_text == VERSION_TABLE_FLAG:
                activeTable.cell(1, 0).paragraphs[0].text = new_dict['[currentDate]']
                activeTable.cell(1, 1).paragraphs[0].text = f"v1.{new_dict['[projectVersion]']}"
                activeTable.cell(1, 2).paragraphs[0].text = new_dict['[projectTitle]']
                activeTable.cell(1, 3).paragraphs[0].text = new_dict['[currentCommercialName]']

            if first_row_column_text == PROJECT_DURATION_TABLE_FLAG:
                activeTable.cell(0, 1).paragraphs[0].text = f"{new_dict['[projectDuration]']} días laborables"

            if first_row_column_text in SECTIONS_WITH_TABLES.keys() and SECTIONS_WITH_TABLES[
                first_row_column_text] in secondary_sections_to_delete:
                activeTable._element.getparent().remove(activeTable._element)

        if len(primary_sections_to_delete) > 0:
            delete_sections(doc.paragraphs, 'Heading 1', primary_sections_to_delete)

        if len(secondary_sections_to_delete) > 0:
            delete_sections(doc.paragraphs, 'Heading 2', secondary_sections_to_delete)

        for paragraph in doc.paragraphs:
            for key, value in new_dict.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value.upper())
                    paragraph.style = paragraph.style.name

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        filename = f"{str(new_dict['[currentDate]']).replace('-', '_')}_{str(new_dict['[projectTitle]']).replace(' ', '_')}_v1.{new_dict['[projectVersion]']}.docx"

        return file_stream, filename
